import json
from re import L
from docx.opc.oxml import parse_xml
import canvas
import os
import time
from docx import Document
from tika import parser
import groupdocs_conversion_cloud
from shutil import copyfile



data_path = "POCs/poc-reading-documents/downloads"
tmp_data_path = "POCs/poc-reading-documents/downloads/17661/BeehiveBeeresearch/beehive_research-1.docx"

test = Document(tmp_data_path)

student_data = {}


def parse_docx(file):
    submission = {}
    fulltext = []
    parsed = Document(file)
    for para in parsed.paragraphs:
        fulltext.append(para.text.encode().decode("utf-8", "replace"))
    submission['text'] = '\n'.join(fulltext)
    return submission

def parse_pdf(file):
    submission = {}
    raw = parser.from_file(file)
    submission['text'] = raw['content']
    return submission

def parse_txt(file):
    submission = {}
    with open(file, 'r') as result:
        submission['text'] = result.readlines()
    return submission

def parse_pptx(file):
    submission = {}
    raw = parser.from_file(file)
    submission['text'] = raw['content']
    return submission

def parse_file(submission):
    if submission.endswith(".docx"):
        result = parse_docx(submission)
    elif submission.endswith(".pdf"):
        result = parse_pdf(submission)
    elif submission.endswith(".pptx"):
        result = parse_pptx(submission)
    elif submission.endswith(".txt"):
        result = parse_txt(submission)
    else:
        return None
    return result



if __name__ == "__main__":
    tmp = [name for name in os.listdir(data_path) if os.path.isdir(os.path.join(data_path+ "/{}".format(name)))]
    for student_folder in tmp:
        student_data[student_folder] = {}
        hand_ins_folder = os.path.join(data_path+ "/{}".format(student_folder))
        for assignment in os.listdir(hand_ins_folder):
            student_data[student_folder][assignment] = {}
            submissions = []
            # TODO Fix the hand_ins column to append to the greater dictionary
            hand_ins = [dname for dname in os.listdir(hand_ins_folder + "/{}".format(assignment))]    
            for submission in hand_ins:
                file_submission = hand_ins_folder + "/{assignment}/{submission}".format(assignment=assignment, submission=submission)
                result = None
                if(os.path.isfile(file_submission)):
                    submissions.append(parse_file(file_submission))
            student_data[student_folder][assignment]["hand-ins"] = submissions


    """ Generating JSON file with all student data """
    with open("test.json", "w+") as outfile:
        json.dump(student_data, outfile)


